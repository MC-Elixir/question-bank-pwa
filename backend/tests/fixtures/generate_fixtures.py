"""生成 3 份样本文档."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


def _add_heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def generate_standard(path: Path) -> None:
    doc = Document()
    _add_heading(doc, "第一章 基础知识")
    doc.add_paragraph("1. Python 是一种解释型语言。")
    doc.add_paragraph("A. 对")
    doc.add_paragraph("B. 错")
    doc.add_paragraph("答案：A")
    doc.add_paragraph("解析：Python 由解释器执行。")

    doc.add_paragraph("2. 下列哪个是 Python 的包管理工具？")
    doc.add_paragraph("A. npm")
    doc.add_paragraph("B. pip")
    doc.add_paragraph("C. gem")
    doc.add_paragraph("D. cargo")
    doc.add_paragraph("正确答案 B")
    doc.add_paragraph("答案解析：pip 是 Python 常用包管理工具。")

    doc.add_paragraph("3. 以下哪些是可变类型？")
    doc.add_paragraph("A. list")
    doc.add_paragraph("B. tuple")
    doc.add_paragraph("C. dict")
    doc.add_paragraph("D. str")
    doc.add_paragraph("参考答案：【A、C】")
    doc.add_paragraph("说明：list 与 dict 可变，tuple 与 str 不可变。")
    doc.save(path)


def generate_missing_answer(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("一、异常样本")
    doc.add_paragraph("1. 地球是行星。")
    doc.add_paragraph("A. 对")
    doc.add_paragraph("B. 错")
    # 故意缺答案
    doc.add_paragraph("解析：这是缺答案样本。")

    doc.add_paragraph("2. 选择首都")
    doc.add_paragraph("A. 上海")
    doc.add_paragraph("B. 北京")
    doc.add_paragraph("C. 广州")
    # 故意缺答案
    doc.save(path)


def generate_mixed_format(path: Path) -> None:
    doc = Document()
    _add_heading(doc, "第二章 混合格式")
    doc.add_paragraph("（1）HTTP 默认端口是 80。")
    doc.add_paragraph("（A）对")
    doc.add_paragraph("（B）错")
    doc.add_paragraph("答案：对")

    doc.add_paragraph("第 2 题 下列协议工作在传输层的是？")
    doc.add_paragraph("A、TCP")
    doc.add_paragraph("B、IP")
    doc.add_paragraph("C、HTTP")
    doc.add_paragraph("D、Ethernet")
    doc.add_paragraph("答案：A")

    doc.add_paragraph("3、可多选：常见数据库有哪些？")
    doc.add_paragraph("A：MySQL")
    doc.add_paragraph("B：Redis")
    doc.add_paragraph("C：Nginx")
    doc.add_paragraph("D：PostgreSQL")
    doc.add_paragraph("答案：A、B、D")
    doc.add_paragraph("解析：Nginx 是 Web 服务器，不是数据库。")
    doc.save(path)


def main() -> None:
    out = Path(__file__).resolve().parent
    out.mkdir(parents=True, exist_ok=True)
    generate_standard(out / "standard_mixed.docx")
    generate_missing_answer(out / "missing_answer.docx")
    generate_mixed_format(out / "mixed_format.docx")
    print("fixtures generated in", out)


if __name__ == "__main__":
    main()
